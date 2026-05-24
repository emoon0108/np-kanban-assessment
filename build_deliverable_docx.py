from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "ethan_moon_task_manager_assessment.docx"


def set_run(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    set_run(run, 16 if level == 1 else 13, True, "2E74B5" if level < 3 else "1F4D78")
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(item)
        set_run(run, 11)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(item)
        set_run(run, 11)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run(run, 11)


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
run = title.add_run("Ethan Moon Task Manager Assessment")
set_run(run, 24, True, "172033")

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run("NextPlay Task Board | Software Development Internship Assessment")
set_run(run, 11, False, "667085")

add_heading(doc, "Overview")
add_body(
    doc,
    "NextPlay Task Board is a responsive Kanban-style task manager built as a static web application with Supabase-backed guest persistence. The interface uses a clean operational layout inspired by Linear and Asana: a compact workspace sidebar, summary stats, clear column hierarchy, and task cards designed for fast scanning.",
)

add_heading(doc, "Links")
add_bullets(
    doc,
    [
        "Live frontend app: TODO - paste the deployed Vercel, Netlify, or Cloudflare Pages URL.",
        "GitHub repository: TODO - paste the public or shared private repository URL.",
    ],
)

add_heading(doc, "Design Decisions")
add_body(
    doc,
    "The board is intentionally dense and work-focused rather than marketing-oriented. Columns stay visually distinct through status color dots, count badges, and contained drop zones. Cards prioritize the task title, priority, labels, due date, and assignee so users can understand the board at a glance.",
)
add_body(
    doc,
    "The app includes responsive behavior for smaller screens by preserving horizontal board scanning while stacking controls and sidebar content.",
)

add_heading(doc, "Advanced Features Built")
add_bullets(
    doc,
    [
        "Team members and assignees with color-coded avatars.",
        "Labels/tags shown on cards and included in search.",
        "Task comments with timestamps inside each task detail modal.",
        "Task activity log for creation, movement, assignment, priority, title, due-date changes, and comments.",
        "Due date indicators for due-soon and overdue tasks.",
        "Search across task content plus priority, assignee, and label filtering.",
        "Board summary stats for total, completed, and overdue tasks.",
        "Local fallback mode for UI review before Supabase credentials are configured.",
    ],
)

add_heading(doc, "Local Setup")
add_numbered(
    doc,
    [
        "Clone the GitHub repository.",
        "Create a free Supabase project.",
        "Enable anonymous sign-ins in Supabase Auth.",
        "Run the schema SQL in the Supabase SQL editor.",
        "Add the Supabase project URL and public anon key at the top of app.js.",
        "Start a local static server with python3 -m http.server 4173.",
        "Open http://localhost:4173.",
    ],
)

add_heading(doc, "Database Schema")
schema = (ROOT / "schema.sql").read_text()
for chunk in [schema[i : i + 1200] for i in range(0, len(schema), 1200)]:
    add_code(doc, chunk)

add_heading(doc, "Tradeoffs and Future Improvements")
add_body(
    doc,
    "The app calls Supabase directly from the frontend, which keeps the project simple and appropriate for the assessment scope. With more time, I would add multi-assignee support, editable/deletable comments, optimistic update rollback for failed drops, keyboard-accessible drag alternatives, and automated end-to-end tests across two anonymous guest sessions.",
)

doc.save(OUT)
print(OUT)
